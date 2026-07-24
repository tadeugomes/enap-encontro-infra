# -*- coding: utf-8 -*-
"""
Gera o artigo cientifico do projeto ENAP Encontro Infra em formato .docx.

Formatacao: Times New Roman 12, espacamento interlinear 1,5, recuo de
primeira linha de 1,25 cm. Resumo com ate 200 palavras.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
import os

# Tipos de relacionamento/conteudo para a parte de notas de rodape
FOOTNOTES_CT = ("application/vnd.openxmlformats-officedocument."
                "wordprocessingml.footnotes+xml")
FOOTNOTES_RT = ("http://schemas.openxmlformats.org/officeDocument/2006/"
                "relationships/footnotes")


def setup_footnotes(document, notes):
    """Cria a parte de notas de rodape (footnotes.xml) com varias notas.

    notes: lista de textos; os ids atribuidos sao 1, 2, ... na ordem da lista.
    """
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    blocos = [
        '<w:footnote w:type="separator" w:id="-1">'
        '<w:p><w:r><w:separator/></w:r></w:p></w:footnote>',
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>',
    ]
    for i, text in enumerate(notes, start=1):
        blocos.append(
            '<w:footnote w:id="%(id)d"><w:p>'
            '<w:r><w:rPr><w:vertAlign w:val="superscript"/>'
            '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            '</w:rPr><w:footnoteRef/></w:r>'
            '<w:r><w:rPr>'
            '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            '<w:sz w:val="20"/></w:rPr>'
            '<w:t xml:space="preserve"> %(text)s</w:t></w:r>'
            '</w:p></w:footnote>' % {"id": i, "text": text}
        )
    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="%s">%s</w:footnotes>' % (w, "".join(blocos))
    )
    partname = PackURI("/word/footnotes.xml")
    part = Part(partname, FOOTNOTES_CT,
                footnotes_xml.encode("utf-8"), document.part.package)
    document.part.relate_to(part, FOOTNOTES_RT)


def add_footnote_ref(paragraph, fid):
    """Anexa um marcador sobrescrito de nota de rodape ao paragrafo."""
    run = paragraph.add_run()
    rpr = run._element.get_or_add_rPr()
    valign = OxmlElement("w:vertAlign")
    valign.set(qn("w:val"), "superscript")
    rpr.append(valign)
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), str(fid))
    run._element.append(ref)

# ---------------------------------------------------------------------------
# Documento e estilo base
# ---------------------------------------------------------------------------
doc = Document()

sec = doc.sections[0]
sec.top_margin = Cm(3)
sec.left_margin = Cm(3)
sec.bottom_margin = Cm(2)
sec.right_margin = Cm(2)

normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn("w:ascii"), "Times New Roman")
rfonts.set(qn("w:hAnsi"), "Times New Roman")
rfonts.set(qn("w:cs"), "Times New Roman")
pf = normal.paragraph_format
pf.line_spacing = 1.5
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.first_line_indent = Cm(1.25)
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")


def add_par(text, indent=True, align="justify", bold=False, italic=False,
            size=12, space_after=0, space_before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25) if indent else Cm(0)
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.alignment = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.keep_with_next = True
    r = p.add_run(text)
    _set_run_font(r, size=12, bold=True)
    return p


def add_table(headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.paragraph_format.first_line_indent = Cm(0)
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after = Pt(2)
        cap.paragraph_format.line_spacing = 1.0
        rc = cap.add_run(caption)
        _set_run_font(rc, size=11, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].paragraphs[0].paragraph_format.first_line_indent = Cm(0)
        hdr[i].paragraphs[0].paragraph_format.line_spacing = 1.0
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hdr[i].paragraphs[0].add_run(h)
        _set_run_font(run, size=11, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            para = cells[i].paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.line_spacing = 1.0
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(str(val))
            _set_run_font(run, size=11)
    return table


def add_source(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    _set_run_font(r, size=10)


# ===========================================================================
# CONTEUDO
# ===========================================================================

# Registra as duas notas de rodape (ids 1 e 2)
setup_footnotes(doc, [
    "Repositório do projeto disponível em: "
    "https://github.com/tadeugomes/enap-encontro-infra.",
    "Aplicação de demonstração do modelo preditivo de severidade disponível "
    "em: https://github.com/fabiokp/defesa_civil_deploy.",
])

# --- Titulo ---
titulo_par = add_par(
    "INTELIGÊNCIA ARTIFICIAL NA GESTÃO DE RECURSOS DE RECONSTRUÇÃO "
    "PÓS-DESASTRE: CAPACIDADE MUNICIPAL, COMPORTAMENTO DE DEMANDA E "
    "AUDITORIA AUTOMATIZADA DE GASTOS PÚBLICOS NO BRASIL (2017-2025)",
    indent=False, align="center", bold=True, space_after=12,
)
# Nota de rodape 1 (repositorio do projeto) ancorada ao titulo
add_footnote_ref(titulo_par, 1)

add_par(
    "Tadeu Gomes Teixeira; Fábio Paim; Alexandre Galardinovic; Bruno Gustavo; "
    "Rafael Giacomin; Rafael Prata; Thiago Mello",
    indent=False, align="center", size=11, space_after=12,
)

# --- Resumo (<= 200 palavras) ---
add_heading("Resumo")
resumo = (
    "Entre 2013 e 2024, os desastres provocaram cerca de R$ 732 bilhões em "
    "prejuízos no Brasil e atingiram 95% dos municípios, em um contexto de "
    "marcada heterogeneidade de capacidades estatais subnacionais. O objetivo "
    "deste artigo é relatar um projeto de inteligência de dados, desenvolvido "
    "no âmbito da ENAP, que articula duas soluções complementares: um motor de "
    "previsão da severidade dos desastres e um sistema de inteligência para a "
    "reconstrução, voltado à alocação de recursos e à auditoria de gastos. A "
    "metodologia mobiliza técnicas de aprendizado de máquina (classificação, "
    "agrupamento, regressão quantílica e detecção de anomalias), aplicadas a "
    "registros oficiais, entre os quais 6.385 processos de reconstrução "
    "(2017-2025), que somam R$ 27,68 bilhões. Os resultados indicam que o tipo "
    "de desastre e a localização são os principais preditores de severidade; "
    "que municípios de baixa capacidade concentram 51% dos recursos e demandam "
    "valor médio 2,94 vezes superior ao dos de alta capacidade; que o "
    "comportamento real de demanda supera o rótulo oficial de capacidade; e "
    "que um modelo de triagem identificou 580 processos com valores "
    "potencialmente superdimensionados, evidenciando o potencial e os limites "
    "da auditoria algorítmica no setor público."
)
# Quebra o resumo na mencao ao motor de previsao para ancorar a nota 2 (deploy)
_marcador = "um motor de previsão da severidade dos desastres"
_idx = resumo.index(_marcador) + len(_marcador)
_resumo_p = add_par(resumo[:_idx], indent=False)
add_footnote_ref(_resumo_p, 2)
_run_resto = _resumo_p.add_run(resumo[_idx:])
_set_run_font(_run_resto, size=12)

add_par(
    "Palavras-chave: capacidade estatal municipal; gestão de desastres; "
    "aprendizado de máquina; previsão de desastres; auditoria de gastos públicos.",
    indent=False, space_before=6,
)

# ===========================================================================
add_heading("1 Introdução")

intro = [
    "Os desastres associados a eventos climáticos extremos deixaram de "
    "constituir ocorrências excepcionais no Brasil e passaram a integrar a "
    "agenda permanente da gestão pública. Inundações, enxurradas, "
    "deslizamentos e tempestades convectivas recorrentes impõem, ano após "
    "ano, a necessidade de mobilizar recursos federais para a reconstrução de "
    "infraestrutura urbana e de serviços essenciais em municípios atingidos. "
    "Entre 2017 e 2025, os registros analisados neste estudo somam 6.385 "
    "processos de reconstrução e R$ 27,68 bilhões em valores solicitados, "
    "distribuídos por 2.113 municípios e 26 unidades da federação. A "
    "magnitude desses números revela que a política de reconstrução "
    "pós-desastre opera em escala que exige instrumentos de gestão capazes de "
    "conciliar celeridade na resposta, equidade na alocação e controle sobre "
    "a aplicação dos recursos. Trata-se, portanto, de um campo em que as "
    "decisões administrativas produzem efeitos imediatos sobre populações "
    "fragilizadas, o que torna especialmente sensível qualquer ineficiência ou "
    "irregularidade na destinação dos valores.",

    "A dimensão do problema, no plano nacional, é expressiva. Segundo o "
    "Panorama dos Desastres no Brasil, da Confederação Nacional de Municípios, "
    "entre 2013 e 2024 os desastres provocaram cerca de R$ 732,2 bilhões em "
    "prejuízos econômicos e atingiram 95% dos municípios brasileiros, que "
    "registraram 70.361 decretos de situação de emergência ou de calamidade "
    "pública (CONFEDERAÇÃO NACIONAL DE MUNICÍPIOS, 2025). No mesmo período, as "
    "decretações de anormalidade cresceram 64%, mesmo quando excluídos os "
    "registros associados à pandemia de Covid-19. O impacto humano acompanha "
    "essa escala: foram contabilizados 2.978 óbitos, mais de um milhão de "
    "desabrigados e cerca de cinco milhões de desalojados, com os danos "
    "econômicos concentrados em dois extremos climáticos, a seca e o excesso "
    "de chuvas, que respondem, somados, por 68,9% das decretações de "
    "anormalidade no país.",

    "A despeito da existência de uma política nacional estruturada, persiste "
    "um abismo entre o dano registrado e a ação executada. Dos R$ 732,2 "
    "bilhões em prejuízos, a União repassou cerca de R$ 5,3 bilhões para "
    "ações de gestão de risco e prevenção de desastres, montante que "
    "corresponde a apenas 39,8% do orçamento federal autorizado para esse fim "
    "no período. A esse descompasso "
    "soma-se um problema de informação: somente 48,3% dos municípios que "
    "decretaram anormalidade informaram os valores dos prejuízos ao sistema "
    "federal, o que indica que os custos reais são significativamente "
    "subdimensionados. Diante desse quadro, a qualificação da alocação e do "
    "controle dos recursos, bem como a antecipação da severidade dos eventos, "
    "constituem desafios cuja superação depende, em larga medida, da "
    "capacidade de extrair inteligência dos dados disponíveis.",

    "A capacidade de um município responder a um desastre, contudo, não se "
    "distribui de maneira uniforme pelo território nacional. A literatura "
    "sobre capacidades estatais subnacionais tem demonstrado que os "
    "municípios brasileiros apresentam dotações muito desiguais de recursos "
    "administrativos, fiscais e técnicos, o que condiciona sua aptidão para "
    "formular projetos, executar obras e prestar contas (GRIN; ABRUCIO, 2018; "
    "MARENCO; STROHSCHOEN; JONER, 2017). Essa heterogeneidade é relevante "
    "porque os mecanismos de transferência voluntária de recursos pressupõem, "
    "em alguma medida, que o ente subnacional disponha de estrutura para "
    "acessar e operar os instrumentos disponíveis. Nesse sentido, a "
    "desigualdade de capacidades pode reproduzir, no momento da reconstrução, "
    "as mesmas assimetrias que tornaram certos municípios mais vulneráveis, de "
    "modo que o desastre opera menos como um evento puramente natural e mais "
    "como o desfecho de processos sociais e institucionais de longa duração.",

    "O reconhecimento dessa desigualdade motivou a construção de indicadores "
    "sintéticos de capacidade institucional, entre os quais o Índice de "
    "Capacidade Municipal (ICM), que classifica os municípios em faixas e "
    "busca orientar a focalização de políticas de fortalecimento. A pergunta "
    "que organiza este artigo emerge do cruzamento entre esse indicador e os "
    "dados de demanda por reconstrução: por que municípios classificados em "
    "faixas de menor capacidade demandam, em média, volumes de recursos "
    "superiores aos solicitados por municípios de maior capacidade? A questão "
    "não é trivial, pois contraria a intuição de que maior capacidade "
    "institucional estaria associada a maior acesso a recursos, e abre espaço "
    "para hipóteses concorrentes sobre vulnerabilidade física, exposição a "
    "eventos severos e qualidade dos dados administrativos. Cada uma dessas "
    "hipóteses remete a um arranjo distinto de causas e, por consequência, a "
    "diferentes recomendações de política, o que confere à pergunta relevância "
    "que ultrapassa o interesse descritivo.",

    "Diante disso, o objetivo geral deste trabalho é analisar a relação entre "
    "a capacidade institucional municipal e o comportamento de demanda por "
    "recursos de reconstrução, bem como avaliar a viabilidade de utilizar "
    "técnicas de aprendizado de máquina para apoiar a auditoria desses gastos. "
    "Como objetivos específicos, busca-se: caracterizar a distribuição dos "
    "valores solicitados segundo as faixas de capacidade; segmentar os "
    "municípios a partir do comportamento efetivo de demanda, e não apenas do "
    "rótulo oficial; estimar a probabilidade de aprovação dos processos e "
    "identificar seus determinantes; e construir um mecanismo de triagem "
    "capaz de sinalizar valores potencialmente superdimensionados. Trata-se, "
    "portanto, de articular um diagnóstico empírico a uma proposta "
    "instrumental de governança, de maneira que a análise dos dados não se "
    "encerre na descrição do fenômeno, mas se traduza em ferramentas "
    "utilizáveis pelos gestores responsáveis pela política.",

    "Cabe esclarecer que esses objetivos integram um projeto mais amplo, "
    "desenvolvido no âmbito do Laboratório de Políticas Públicas Sociais da "
    "ENAP, que articula duas soluções complementares de inteligência de dados. "
    "A primeira, um motor de previsão da severidade dos desastres, atua nas "
    "fases de prevenção e preparação, ao antecipar a magnitude provável dos "
    "danos de um novo evento. A segunda, um sistema de inteligência para a "
    "reconstrução, atua na fase de recuperação, ao otimizar a alocação de "
    "recursos e auditar os gastos. Embora o eixo central deste artigo seja a "
    "segunda solução, a primeira é igualmente apresentada, por compor, com "
    "ela, um ecossistema integrado que percorre o ciclo de gestão de riscos "
    "estabelecido pela Política Nacional de Proteção e Defesa Civil.",

    "A justificativa do estudo assenta-se em duas ordens de consideração. Do "
    "ponto de vista acadêmico, a articulação entre a literatura de capacidades "
    "estatais e as técnicas de ciência de dados ainda é incipiente no Brasil, "
    "de modo que a aplicação desse instrumental a um problema concreto de "
    "política pública contribui para um campo em formação. Do ponto de vista "
    "prático, o volume de recursos envolvidos e a recorrência crescente dos "
    "desastres tornam imperativo dispor de mecanismos capazes de qualificar "
    "tanto a alocação quanto o controle dos gastos. A análise individualizada "
    "de milhares de processos excede a capacidade operacional dos órgãos de "
    "controle, o que recoloca, em novos termos, a velha tensão entre a "
    "amplitude da fiscalização e a escassez dos meios disponíveis para "
    "exercê-la.",

    "O argumento desenvolvido ao longo do texto sustenta que os indicadores "
    "formais de capacidade, embora úteis para a focalização de políticas, "
    "capturam de modo apenas parcial o comportamento real dos municípios "
    "diante dos desastres, e que a observação desse comportamento por meio de "
    "dados transacionais oferece subsídios mais precisos tanto para a alocação "
    "quanto para o controle. Esse argumento será sustentado pela convergência "
    "de resultados de fases analíticas distintas, que apontam, de maneira "
    "independente, para a limitação explicativa do rótulo oficial quando "
    "confrontado com o comportamento observado. Cabe registrar que o estudo se "
    "apoia em dados administrativos secundários, com as limitações inerentes a "
    "essa natureza, e que os modelos aqui apresentados são concebidos como "
    "instrumentos de apoio à decisão, e não como substitutos do julgamento "
    "técnico.",

    "O artigo está organizado em seis seções, além desta introdução. A segunda "
    "seção apresenta o referencial teórico, que articula a discussão sobre "
    "capacidades estatais municipais e desastres, a literatura sobre "
    "governança e controle de gastos públicos e os fundamentos das técnicas de "
    "aprendizado de máquina mobilizadas. A terceira seção descreve a "
    "metodologia, detalhando as fontes de dados, os procedimentos de "
    "tratamento e o desenho analítico em cinco fases. A quarta seção expõe os "
    "resultados, organizados segundo os objetivos específicos. A quinta seção "
    "discute os achados à luz do referencial teórico e explicita as limitações "
    "do estudo. A sexta seção, por fim, sintetiza as contribuições e propõe "
    "uma agenda de pesquisa.",
]
for p in intro:
    add_par(p)

# ===========================================================================
add_heading("2 Referencial teórico")

add_heading("2.1 Capacidade estatal municipal e desastres", level=2)
ref1 = [
    "O conceito de capacidade estatal tem sido mobilizado para descrever a "
    "aptidão das organizações governamentais para implementar suas decisões e "
    "alcançar os objetivos que se propõem. No plano subnacional, essa noção "
    "ganha contornos específicos, pois os municípios brasileiros acumularam, "
    "a partir da Constituição de 1988, um amplo conjunto de competências sem "
    "que houvesse correspondente equalização de recursos administrativos e "
    "fiscais (ABRUCIO, 2005). Verifica-se, assim, um descompasso entre a "
    "responsabilidade atribuída ao ente local e os meios de que dispõe para "
    "exercê-la, descompasso que se manifesta com particular nitidez em "
    "situações de emergência, quando a demanda por respostas rápidas e "
    "tecnicamente qualificadas se intensifica. Em tais momentos, a fragilidade "
    "das estruturas locais não apenas retarda a resposta, mas pode "
    "comprometer a própria capacidade de acessar os recursos federais "
    "disponíveis, gerando um efeito que aprofunda as desigualdades preexistentes.",

    "A literatura recente sobre capacidades estatais municipais tem buscado "
    "mensurar essa heterogeneidade por meio de indicadores que combinam "
    "dimensões administrativas, fiscais e de pessoal (MARENCO; STROHSCHOEN; "
    "JONER, 2017; GRIN; ABRUCIO, 2018). Esses estudos evidenciam que a "
    "capacidade se concentra em municípios de maior porte e renda, ao passo "
    "que os municípios menores e mais pobres, frequentemente os mais expostos "
    "a riscos, dispõem de estruturas mais frágeis. Dessa forma, a capacidade "
    "institucional opera como variável mediadora entre a ocorrência de um "
    "evento adverso e a efetividade da resposta pública, condicionando desde "
    "a elaboração de planos de trabalho até a prestação de contas. O Índice de "
    "Capacidade Municipal utilizado neste estudo inscreve-se nessa tradição, "
    "ao sintetizar múltiplas dimensões em uma classificação por faixas que "
    "pretende ordenar os municípios segundo sua aptidão institucional.",

    "No campo específico da gestão de desastres, a Política Nacional de "
    "Proteção e Defesa Civil, instituída pela Lei nº 12.608, de 2012, "
    "estabeleceu diretrizes para a atuação integrada dos entes federativos em "
    "ações de prevenção, mitigação, preparação, resposta e recuperação. A "
    "norma consagrou a perspectiva de gestão de riscos, deslocando o foco da "
    "resposta emergencial para o ciclo completo que inclui a prevenção, e "
    "atribuiu responsabilidades específicas a cada nível de governo. Em "
    "complemento, o Plano Nacional de Adaptação à Mudança do Clima incorporou "
    "a gestão de risco de desastres como estratégia setorial, ao reconhecer a "
    "relação entre a intensificação dos eventos extremos e as mudanças "
    "climáticas (BRASIL, 2016). A "
    "literatura sociológica sobre o tema, por sua vez, tem insistido em que os "
    "desastres não constituem fenômenos puramente naturais, mas resultam da "
    "interação entre ameaças e vulnerabilidades socialmente produzidas "
    "(VALENCIO, 2009). Sob essa perspectiva, a maior demanda por recursos "
    "observada em municípios vulneráveis pode ser lida como expressão de "
    "vulnerabilidades acumuladas, e não como mera consequência da intensidade "
    "física dos eventos.",

    "Essa distinção entre ameaça e vulnerabilidade é central para a "
    "interpretação dos resultados que se seguem. Se o dano decorresse apenas da "
    "intensidade física do evento, seria de esperar que a demanda por recursos "
    "se distribuísse de modo relativamente independente da capacidade "
    "institucional. A constatação de uma associação sistemática entre baixa "
    "capacidade e maior demanda sugere, ao contrário, que a vulnerabilidade "
    "social e institucional desempenha papel ativo na determinação dos danos. "
    "Diante disso, o estudo da reconstrução pós-desastre não pode prescindir de "
    "uma leitura que articule as condições materiais de cada município à sua "
    "trajetória institucional, sob pena de reduzir um fenômeno social complexo "
    "a uma questão meramente técnica de dimensionamento de obras.",

    "A organização federativa brasileira acrescenta complexidade ao quadro. As "
    "transferências de recursos para reconstrução articulam competências da "
    "União, dos estados e dos municípios, em um arranjo no qual a "
    "responsabilidade pela execução recai, em última instância, sobre o ente "
    "local, ainda que os recursos e os critérios de elegibilidade sejam "
    "definidos em âmbito federal. Esse desenho pressupõe que o município "
    "disponha de burocracia capaz de traduzir a necessidade local na linguagem "
    "técnica exigida pelos instrumentos de transferência, o que recoloca a "
    "questão da capacidade estatal no centro da política. Diante disso, a "
    "análise da demanda por reconstrução não pode dissociar-se das condições "
    "institucionais em que essa demanda é formulada, sob risco de atribuir ao "
    "comportamento do município características que, na verdade, decorrem das "
    "regras e dos constrangimentos do próprio arranjo federativo.",
]
for p in ref1:
    add_par(p)

add_heading("2.2 Governança, transparência e controle de gastos públicos", level=2)
ref2 = [
    "A alocação de recursos para reconstrução insere-se em um arranjo de "
    "governança que envolve a União, os estados e os municípios, e que "
    "mobiliza instrumentos de transferência condicionados à apresentação de "
    "planos de trabalho e à comprovação da aplicação dos valores. Nesse "
    "arranjo, a transparência e o controle assumem papel central, tanto para "
    "assegurar a regularidade dos gastos quanto para sustentar a confiança "
    "pública na política. A literatura sobre controle da administração pública "
    "tem destacado a complementaridade entre os controles interno, externo e "
    "social, bem como a importância de mecanismos que permitam identificar "
    "irregularidades sem comprometer a celeridade da resposta (SPECK, 2002). "
    "Essa tensão entre controle e celeridade é particularmente aguda no caso "
    "dos desastres, em que a demora na liberação dos recursos pode agravar o "
    "sofrimento das populações atingidas, mas a liberação sem o devido "
    "escrutínio abre margem para o desperdício e a apropriação indevida.",

    "O exame de grandes volumes de processos, contudo, esbarra em limites "
    "práticos do controle tradicional, fundado na análise individualizada de "
    "cada solicitação. Diante de milhares de processos e bilhões de reais, a "
    "capacidade de escrutínio dos órgãos de controle é necessariamente "
    "seletiva, o que reforça a relevância de instrumentos capazes de priorizar "
    "casos a partir de critérios objetivos. Sendo assim, a detecção de "
    "anomalias em dados de gasto público desponta como estratégia para "
    "direcionar a atenção dos auditores aos processos que mais se afastam dos "
    "padrões esperados, otimizando o uso de recursos escassos de fiscalização. "
    "Essa abordagem não substitui o controle convencional, mas o qualifica, ao "
    "permitir que o esforço humano de verificação se concentre nos casos de "
    "maior risco potencial, em vez de dispersar-se de modo uniforme por todo o "
    "universo de processos.",

    "Importa, no entanto, situar os limites dessa estratégia. A sinalização de "
    "um processo como atípico não equivale a um juízo de irregularidade, pois "
    "o desvio em relação ao padrão pode decorrer de circunstâncias legítimas, "
    "de erros de registro ou, de fato, de práticas irregulares. A confusão "
    "entre essas possibilidades pode produzir tanto a estigmatização indevida "
    "de gestores honestos quanto a falsa sensação de que o controle algorítmico "
    "dispensa o julgamento técnico. Dessa maneira, a incorporação de "
    "ferramentas automatizadas ao controle de gastos exige que se preserve, em "
    "todas as etapas, o caráter de apoio à decisão, reservando-se ao auditor "
    "humano a competência para distinguir, no exame dos casos priorizados, "
    "entre o que é apenas incomum e o que é efetivamente irregular.",
]
for p in ref2:
    add_par(p)

add_heading("2.3 Aprendizado de máquina aplicado ao setor público", level=2)
ref3 = [
    "A incorporação de técnicas de aprendizado de máquina à gestão pública tem "
    "se expandido, com aplicações que vão da previsão de demanda por serviços "
    "à detecção de fraudes em transações governamentais. Os algoritmos "
    "mobilizados neste estudo pertencem a famílias consolidadas na literatura "
    "estatística e computacional. O agrupamento por k-médias, formalizado por "
    "MacQueen (1967), permite segmentar observações em grupos internamente "
    "homogêneos sem rótulos prévios, sendo particularmente útil quando se "
    "deseja descobrir padrões latentes em dados não rotulados. As florestas "
    "aleatórias, propostas por Breiman (2001), combinam múltiplas árvores de "
    "decisão para produzir classificações robustas e estimativas de "
    "importância de variáveis, característica que as torna especialmente "
    "valiosas em contextos nos quais a interpretabilidade dos determinantes "
    "importa tanto quanto a acurácia da predição.",

    "Para os objetivos de auditoria, duas técnicas merecem destaque. A "
    "regressão quantílica, introduzida por Koenker e Bassett (1978), estima "
    "diferentes quantis da distribuição condicional da variável resposta, o "
    "que possibilita construir faixas de valores esperados em vez de uma única "
    "previsão pontual. Essa propriedade é decisiva para a auditoria, pois "
    "permite definir limites inferior e superior dentro dos quais um valor "
    "pode ser considerado compatível com casos semelhantes. O gradient "
    "boosting, sistematizado por Friedman (2001), constrói modelos preditivos "
    "por meio da adição sequencial de estimadores fracos, e pode ser combinado "
    "à perda quantílica para gerar tais faixas. A floresta de isolamento, "
    "proposta por Liu, Ting e Zhou (2008), identifica observações anômalas a "
    "partir da facilidade com que são isoladas em partições aleatórias do "
    "espaço de atributos, oferecendo uma via independente para a detecção de "
    "casos atípicos.",

    "A combinação dessas abordagens permite tratar a auditoria não como um "
    "juízo binário sobre cada processo, mas como um exercício de posicionamento "
    "de cada solicitação em relação a um padrão estimado a partir de casos "
    "comparáveis. Cabe ressalvar, todavia, que a aplicação desses métodos ao "
    "setor público requer cautela interpretativa. A qualidade dos dados "
    "administrativos e a presença de erros de registro podem afetar de modo "
    "expressivo os resultados, e a opacidade de certos modelos pode dificultar "
    "a prestação de contas sobre as decisões por eles apoiadas. Por essa razão, "
    "privilegiaram-se, neste estudo, técnicas que oferecem alguma medida de "
    "interpretabilidade, como a importância de variáveis das florestas "
    "aleatórias e as faixas explícitas da regressão quantílica, de maneira a "
    "permitir que os achados sejam escrutinados e contestados.",

    "No contexto brasileiro, a adoção de técnicas analíticas pela administração "
    "pública tem avançado de modo desigual, concentrando-se em órgãos de "
    "controle e em iniciativas pontuais de modernização da gestão. Experiências "
    "de mineração de dados em tribunais de contas e em órgãos de fiscalização "
    "têm demonstrado o potencial dessas ferramentas para a identificação de "
    "padrões suspeitos em contratações e transferências, ao mesmo tempo em que "
    "evidenciam os desafios de integração de bases de dados fragmentadas e de "
    "qualificação das equipes responsáveis. O presente estudo inscreve-se nesse "
    "movimento, ao aplicar um conjunto de técnicas consolidadas a um domínio "
    "específico de política pública, com a preocupação de articular o rigor "
    "metodológico à utilidade prática para os gestores e os órgãos de controle.",
]
for p in ref3:
    add_par(p)

# ===========================================================================
add_heading("3 Metodologia")

add_heading("3.1 Fontes de dados e tratamento", level=2)
met1 = [
    "O estudo apoia-se em duas bases de dados administrativos. A primeira "
    "corresponde aos relatórios de acompanhamento de processos de "
    "reconstrução, abrangendo o período de 2017 a 2025, e reúne informações "
    "sobre unidade da federação, município, tipo de desastre, valor "
    "solicitado, situação do processo e datas de criação e solicitação. Após "
    "a consolidação, essa base totaliza 6.385 processos, distribuídos por "
    "2.113 municípios, 26 unidades da federação e 29 tipos de desastre, "
    "somando R$ 27,68 bilhões em valores solicitados. A segunda base "
    "corresponde ao Índice de Capacidade Municipal (ICM), que classifica os "
    "municípios em quatro faixas de capacidade institucional, designadas de A "
    "(alta) a D (baixa), e que, após a limpeza, contempla 5.445 municípios. A "
    "natureza secundária e administrativa desses dados impõe que se reconheça, "
    "desde o início, a dependência do estudo em relação à qualidade do "
    "registro original.",

    "O cruzamento das duas bases exigiu um conjunto de procedimentos de "
    "tratamento. A consolidação dos arquivos originais foi seguida da "
    "padronização dos nomes de municípios e da conversão dos valores "
    "monetários, originalmente registrados como texto, para formato numérico. "
    "Constatou-se a presença de 152 municípios duplicados com faixas de "
    "capacidade conflitantes, situação tratada pela adoção do critério de "
    "benefício, segundo o qual, em caso de conflito, atribuiu-se ao município "
    "a melhor faixa identificada. Essa escolha metodológica, embora "
    "conservadora, busca evitar a penalização indevida de municípios em razão "
    "de inconsistências de registro. Após esses procedimentos, o conjunto "
    "integrado alcançou cobertura de 97,7% dos municípios com processos, com "
    "apenas 2,3% sem correspondência no índice de capacidade, o que assegura "
    "ampla representatividade ao cruzamento.",

    "A etapa de engenharia de atributos preparou as variáveis para os modelos. "
    "Os valores monetários, marcados por forte assimetria, foram transformados "
    "pela função logarítmica (log1p), de modo a aproximar sua distribuição da "
    "normalidade e a reduzir o peso desproporcional dos valores extremos. As "
    "variáveis categóricas, como unidade da federação e tipo de desastre, "
    "foram codificadas numericamente, ao passo que a faixa de capacidade "
    "recebeu codificação ordinal, preservando a ordenação entre as faixas. "
    "Para a agregação por município, computaram-se o número de processos, o "
    "valor total e o valor médio, atributos que sintetizam o comportamento de "
    "demanda. A depender do algoritmo, aplicaram-se padronização por escala "
    "robusta, mais adequada na presença de valores atípicos, ou padronização "
    "convencional, escolha orientada pela sensibilidade de cada técnica aos "
    "valores extremos.",
]
for p in met1:
    add_par(p)

add_heading("3.2 Desenho analítico em cinco fases", level=2)
met2 = [
    "O desenho analítico estrutura-se em um pipeline de cinco fases, "
    "concebidas para responder, de modo progressivo, aos objetivos do estudo. "
    "A primeira fase consiste em diagnóstico e regressão exploratória, voltada "
    "a caracterizar a distribuição dos valores por faixa de capacidade e a "
    "testar a hipótese de associação entre vulnerabilidade institucional e "
    "magnitude da demanda. A segunda fase aplica agrupamento por k-médias para "
    "segmentar os municípios segundo o comportamento efetivo de demanda. A "
    "terceira fase emprega um classificador para estimar a probabilidade de "
    "aprovação dos processos e identificar seus determinantes. A quarta fase "
    "recorre à regressão quantílica para estimar faixas de valor esperado e "
    "sinalizar desvios. A quinta fase, complementar, mobiliza a detecção de "
    "anomalias e projeta um simulador de alertas. Essa organização sequencial "
    "permite que cada fase se beneficie dos achados das anteriores, de modo "
    "que os grupos comportamentais identificados na segunda fase, por exemplo, "
    "tornem-se atributos das fases subsequentes.",

    "Na fase de agrupamento, optou-se por quatro grupos, número escolhido para "
    "permitir comparação direta com as quatro faixas do índice de capacidade, "
    "tendo a métrica de silhueta orientado a avaliação da coesão dos grupos. "
    "Os atributos utilizados foram o logaritmo do valor total, o logaritmo do "
    "valor médio e o número de processos por município, sobre os quais se "
    "aplicou padronização por escala robusta. Na fase de classificação, a "
    "variável resposta foi definida de modo binário, distinguindo os processos "
    "com recurso transferido (aprovados) daqueles indeferidos, excluídos ou "
    "sobrestados (reprovados), enquanto os processos ainda em análise foram "
    "excluídos do treinamento. Adotou-se o classificador de florestas "
    "aleatórias, com cem estimadores e profundidade máxima limitada, partição "
    "dos dados em 80% para treino e 20% para teste, estratificação da variável "
    "resposta e ponderação das classes para mitigar o desequilíbrio entre "
    "aprovados e reprovados.",

    "Na fase de regressão preditiva, empregou-se o gradient boosting com perda "
    "quantílica para estimar os quantis P10, P50 e P90 do valor solicitado, "
    "condicionados à unidade da federação, ao tipo de desastre, à faixa de "
    "capacidade, ao grupo comportamental e à faixa populacional. A variável "
    "resposta foi o logaritmo do valor, e a análise restringiu-se aos "
    "processos aprovados com valores superiores a R$ 1.000, de modo a evitar a "
    "contaminação das estimativas por registros residuais. A partir das faixas "
    "estimadas, cada processo foi classificado como normal, quando situado "
    "entre os limites P10 e P90; baixo, quando inferior ao P10; ou alto, "
    "quando superior ao P90, hipótese que sinaliza valor potencialmente "
    "superdimensionado. A floresta de isolamento, por fim, foi parametrizada "
    "para uma contaminação esperada de 5%, com pré-processamento que combinou "
    "padronização e codificação das variáveis categóricas. O quadro a seguir "
    "sintetiza as fases, os algoritmos e as métricas de avaliação.",
]
for p in met2:
    add_par(p)

add_table(
    ["Fase", "Técnica / Algoritmo", "Objetivo", "Métrica principal"],
    [
        ["1", "Regressão exploratória; Isolation Forest", "Diagnóstico e hipóteses", "Estatística descritiva"],
        ["2", "K-médias (k=4)", "Segmentação comportamental", "Silhueta"],
        ["3", "Florestas aleatórias", "Previsão de aprovação", "ROC-AUC = 0,7972"],
        ["4", "Gradient boosting quantílico", "Faixa de “valor justo”", "Cobertura P10-P90 = 76,3%"],
        ["5", "Isolation Forest; simulador", "Detecção de anomalias", "320 anomalias"],
    ],
    caption="Quadro 1 - Síntese do pipeline analítico em cinco fases",
)
add_source("Fonte: elaboração própria.")

add_par(
    "As métricas de avaliação foram escolhidas conforme a natureza de cada "
    "tarefa. Para a classificação, privilegiou-se a área sob a curva ROC, "
    "métrica robusta ao desequilíbrio entre classes e adequada à comparação "
    "entre a capacidade de discriminação do modelo e o acaso. Para a regressão "
    "quantílica, adotou-se a cobertura do intervalo entre os quantis "
    "estimados, que mede a proporção de observações efetivamente contidas na "
    "faixa prevista, complementada pelo coeficiente de determinação no espaço "
    "logarítmico. Para o agrupamento, a métrica de silhueta avaliou a coesão "
    "interna e a separação entre os grupos. Essa diversidade de métricas "
    "reflete a heterogeneidade das tarefas e previne a avaliação inadequada de "
    "um modelo por critérios concebidos para outra finalidade."
)

add_heading("3.3 O motor de previsão de severidade", level=2)
met3 = [
    "A solução de previsão, complementar ao pipeline de reconstrução, apoia-se "
    "em um modelo de aprendizado de máquina treinado com dados históricos do "
    "Atlas Brasileiro de Desastres Naturais (CEPED/UFSC, 2023), de bases do "
    "Instituto Brasileiro de Geografia e Estatística referentes ao produto "
    "interno bruto municipal e à hierarquia urbana (IBGE, 2020, 2023) e de "
    "indicadores de cobertura da atenção básica do Departamento de Informática "
    "do Sistema Único de Saúde (BRASIL, 2024), com o propósito de antecipar o "
    "impacto de um novo evento. O modelo prevê quatro categorias "
    "de dano. A ocorrência de vítimas é tratada como classificação binária, "
    "orientada a maximizar a detecção de eventos com vítimas fatais ou "
    "feridos. Os danos humanos totais, os danos materiais e os prejuízos "
    "econômicos são tratados como classificação multiclasse em três níveis, "
    "designados nenhum, baixo e alto, de modo a estimar, respectivamente, a "
    "magnitude das pessoas afetadas, o nível de destruição de infraestrutura "
    "e o impacto financeiro para os setores público e privado.",

    "A escolha por classificações em níveis, em vez de previsões pontuais, "
    "responde à elevada variabilidade dos danos e à finalidade de triagem da "
    "ferramenta, que se destina a ordenar eventos por gravidade provável, e "
    "não a estimar valores exatos. A solução foi disponibilizada em ambiente "
    "de demonstração de acesso público, o que permite a consulta interativa "
    "das previsões e evidencia a viabilidade técnica de sua incorporação à "
    "rotina dos órgãos de defesa civil.",
]
for p in met3:
    add_par(p)

add_heading("3.4 Reprodutibilidade", level=2)
add_par(
    "Todas as análises foram implementadas em Python, com o uso das "
    "bibliotecas pandas, NumPy e scikit-learn, e com sementes aleatórias "
    "fixas (random_state = 42) em todos os modelos, de modo a assegurar a "
    "reprodutibilidade dos resultados. O código-fonte, os dados de trabalho e "
    "os artefatos de saída encontram-se disponíveis no repositório do projeto, "
    "o que permite a verificação independente dos números aqui reportados. Os "
    "resultados das cinco fases foram auditados por reexecução dos modelos, "
    "tendo sido reproduzidos dentro de margens compatíveis com pequenas "
    "variações entre versões das bibliotecas."
)

# ===========================================================================
add_heading("4 Resultados")

add_heading("4.1 O paradoxo da capacidade institucional", level=2)
res1 = [
    "A análise da distribuição dos valores por faixa de capacidade revela um "
    "padrão que contraria a intuição inicial. Os municípios classificados na "
    "faixa D, de baixa capacidade institucional, apresentam valor médio "
    "solicitado de R$ 28,68 milhões, ao passo que os municípios da faixa A, de "
    "alta capacidade, solicitam, em média, R$ 9,74 milhões. A razão entre os "
    "dois valores é de 2,94, o que significa que os municípios mais "
    "vulneráveis demandam, em média, quase três vezes mais recursos por "
    "processo do que os mais bem estruturados. A faixa D concentra, ademais, "
    "51% do total de recursos solicitados, embora "
    "reúna cerca de um quarto dos municípios analisados. Esse achado, que se "
    "poderia denominar paradoxo da capacidade institucional, constitui o ponto "
    "de partida das análises subsequentes.",
]
for p in res1:
    add_par(p)

add_table(
    ["Faixa", "Capacidade", "Valor médio (R$)", "% do total"],
    [
        ["A", "Alta", "9,74 milhões", "8%"],
        ["B", "Média-alta", "9,79 milhões", "20%"],
        ["C", "Média-baixa", "6,63 milhões", "19%"],
        ["D", "Baixa", "28,68 milhões", "51%"],
    ],
    caption="Tabela 1 - Valor médio solicitado e participação no total por faixa de capacidade",
)
add_source("Fonte: elaboração própria a partir dos dados de acompanhamento e do ICM.")

res1b = [
    "A leitura desses números, contudo, exige cautela analítica. A mediana dos "
    "valores na faixa D situa-se em patamar inferior ao das demais faixas, o "
    "que indica que a média elevada é fortemente influenciada por um pequeno "
    "número de processos de valor muito alto. Verifica-se, portanto, que a "
    "distribuição na faixa D combina grande quantidade de solicitações "
    "modestas com poucos casos extremos, responsáveis por puxar a média para "
    "cima. Esse achado tem dupla implicação: por um lado, sugere que a "
    "vulnerabilidade institucional pode estar associada a eventos de impacto "
    "financeiro elevado; por outro, adverte contra interpretações que tomem a "
    "média como representativa do município típico da faixa, uma vez que os "
    "valores extremos podem refletir tanto desastres genuinamente severos "
    "quanto erros de registro ou superfaturamentos a investigar. A distinção "
    "entre essas possibilidades, retomada nas fases posteriores, é decisiva "
    "para a interpretação adequada do paradoxo.",

    "O exame da série temporal reforça a relevância do tema. O ano de 2024 "
    "concentrou o maior volume de toda a série histórica, com 1.490 processos "
    "e cerca de R$ 8,6 bilhões em valores solicitados, em patamar muito "
    "superior ao dos anos anteriores. Esse pico, ainda que demande análise "
    "adicional, é compatível com a hipótese de intensificação dos eventos "
    "climáticos extremos e com a concentração de ocorrências em determinadas "
    "regiões, com destaque para o Rio Grande do Sul, que figura entre as "
    "unidades da federação com maior número de processos. A concentração "
    "regional e temporal das ocorrências sugere que a política de reconstrução "
    "opera sob pressão crescente e desigual, o que torna ainda mais premente a "
    "disponibilidade de instrumentos capazes de qualificar a alocação dos "
    "recursos.",

    "O paradoxo da capacidade, longe de constituir mera curiosidade "
    "estatística, recoloca uma questão central para o desenho das políticas de "
    "reconstrução. Se os municípios de menor capacidade são precisamente os "
    "que demandam mais recursos, então a política tende a operar de modo "
    "reativo, transferindo volumes elevados a entes que dispõem de menor "
    "estrutura para aplicá-los de maneira eficiente e para prestar contas. "
    "Dessa forma, o achado aponta para a necessidade de articular a política de "
    "reconstrução a uma política de fortalecimento institucional, sob pena de "
    "que os recursos transferidos não se convertam em capacidade efetiva de "
    "resposta e prevenção. Essa articulação, contudo, depende de uma "
    "compreensão mais fina do comportamento dos municípios, que a mera "
    "classificação por faixas não é capaz de oferecer.",

    "A dimensão regional do fenômeno adiciona uma camada interpretativa que "
    "merece atenção. A concentração de processos em determinadas unidades da "
    "federação, com destaque para o Rio Grande do Sul, sugere que os eventos "
    "extremos não se distribuem aleatoriamente pelo território, mas seguem "
    "padrões geográficos e climáticos que tendem a recair, de modo recorrente, "
    "sobre as mesmas regiões. Essa recorrência territorial dialoga com o achado "
    "do grupo de alta frequência, identificado na segmentação comportamental, "
    "e aponta para a existência de áreas que demandam não apenas respostas "
    "reativas a cada desastre, mas estratégias estruturais de prevenção e "
    "adaptação. Diante disso, a política de reconstrução, se considerada de "
    "forma isolada, arrisca-se a operar como um ciclo de transferências "
    "repetidas a territórios cuja vulnerabilidade permanece, no essencial, "
    "inalterada entre um evento e o seguinte.",
]
for p in res1b:
    add_par(p)

add_heading("4.2 Comportamento de demanda e segmentação dos municípios", level=2)
res2 = [
    "A segunda fase buscou superar os limites da classificação oficial ao "
    "agrupar os municípios segundo o comportamento efetivo de demanda. O "
    "agrupamento por k-médias produziu quatro perfis nitidamente distintos. O "
    "primeiro grupo, de baixo impacto, reúne 76% dos municípios e "
    "caracteriza-se por poucos processos e valores reduzidos. O segundo grupo, "
    "de alto custo, agrega 22% dos municípios e combina desastres de elevado "
    "valor com frequência intermediária. O terceiro grupo, de alta frequência, "
    "reúne apenas 1% dos municípios, atingidos de forma repetida ao longo do "
    "período. O quarto grupo isola um único caso extremo, o do município de "
    "Porto Alegre, cujo padrão de demanda destoa de todos os demais e cuja "
    "presença como grupo autônomo ilustra a sensibilidade da técnica à "
    "existência de observações genuinamente excepcionais.",

    "O cruzamento entre os grupos comportamentais e as faixas de capacidade "
    "produziu o achado mais relevante desta fase. Constatou-se que 81% dos "
    "municípios classificados na faixa D pertencem ao grupo de baixo impacto, "
    "ao passo que diversos municípios da faixa A integram o grupo de alto "
    "custo. Dessa forma, evidencia-se que a faixa oficial de capacidade não "
    "determina o comportamento real de demanda: a vulnerabilidade "
    "institucional não implica, necessariamente, demanda crônica por recursos, "
    "tampouco a alta capacidade protege o município da ocorrência de desastres "
    "custosos. Esse descompasso entre o rótulo formal e o comportamento "
    "observado sustenta o argumento de que os dados transacionais oferecem "
    "informação complementar e, em certos aspectos, mais precisa para a "
    "gestão, na medida em que registram o que de fato ocorre, e não apenas o "
    "que se esperaria a partir das características estruturais do município.",
]
for p in res2:
    add_par(p)

add_table(
    ["Grupo", "Perfil", "% dos municípios", "Característica"],
    [
        ["0", "Baixo impacto", "76%", "Poucos processos, valores baixos"],
        ["1", "Alto custo", "22%", "Desastres caros, frequência média"],
        ["3", "Alta frequência", "1%", "Atingidos repetidamente"],
        ["2", "Outlier extremo", "< 0,1%", "Caso único (Porto Alegre)"],
    ],
    caption="Tabela 2 - Perfis comportamentais identificados pelo agrupamento",
)
add_source("Fonte: elaboração própria.")

add_par(
    "A relevância desse descompasso para a gestão pode ser apreendida em dois "
    "planos. No plano da focalização, a constatação de que a maioria dos "
    "municípios de baixa capacidade integra o grupo de baixo impacto sugere "
    "que a vulnerabilidade institucional, por si só, não é um critério "
    "suficiente para antecipar a demanda por recursos. No plano do "
    "monitoramento, a identificação de um pequeno grupo de alta frequência, "
    "composto por municípios atingidos de modo recorrente, indica a existência "
    "de territórios que demandam atenção continuada e que poderiam beneficiar-"
    "se de medidas estruturais de prevenção, em vez de respostas pontuais a "
    "cada novo evento. Sendo assim, a segmentação comportamental oferece um "
    "mapa mais detalhado do problema do que a classificação por faixas, sem "
    "que isso implique descartar o indicador oficial, cuja utilidade reside "
    "justamente na síntese de características estruturais."
)

add_par(
    "O caso de Porto Alegre, isolado como grupo autônomo pelo agrupamento, "
    "ilustra os limites da análise uniforme. A magnitude e a recorrência de "
    "suas solicitações, associadas aos eventos que atingiram a capital gaúcha, "
    "produziram um padrão de demanda sem paralelo no conjunto dos municípios, a "
    "ponto de a técnica de agrupamento o tratar como categoria própria. Esse "
    "resultado evidencia que os grandes centros urbanos podem concentrar, em um "
    "único ente, volumes de demanda comparáveis aos de regiões inteiras, o que "
    "tensiona os critérios uniformes de análise e sugere a conveniência de "
    "tratamentos diferenciados para casos de escala excepcional. Constata-se, "
    "assim, que a heterogeneidade dos municípios brasileiros se manifesta não "
    "apenas entre faixas de capacidade, mas também no interior de cada faixa, o "
    "que reforça a utilidade de abordagens que combinem a classificação "
    "estrutural à observação do comportamento."
)

add_heading("4.3 Determinantes da aprovação dos processos", level=2)
res3 = [
    "A terceira fase estimou a probabilidade de aprovação dos processos por "
    "meio de um classificador de florestas aleatórias. O modelo alcançou área "
    "sob a curva ROC de 0,7972, o que indica capacidade de discriminação da "
    "ordem de 80% entre processos aprovados e reprovados. A revocação para a "
    "classe dos aprovados atingiu 88%, sinalizando que o modelo identifica "
    "corretamente a maior parte dos processos efetivamente aprovados. Esses "
    "resultados, embora não configurem desempenho próximo da perfeição, são "
    "compatíveis com a complexidade do fenômeno e suficientes para sustentar "
    "uma aplicação de triagem, na qual o objetivo não é decidir "
    "automaticamente sobre cada processo, mas ordenar os casos segundo a "
    "probabilidade de aprovação, de modo a apoiar a priorização do trabalho "
    "técnico.",

    "A análise da importância das variáveis expõe os determinantes da "
    "decisão. O valor solicitado respondeu por 63% da importância total, "
    "consolidando-se como o determinante dominante da aprovação, o que sugere "
    "que solicitações muito afastadas do padrão enfrentam escrutínio "
    "diferenciado. O conjunto formado pela faixa populacional e pelo grupo "
    "comportamental respondeu por 15%, a unidade da federação por 10% e o tipo "
    "de desastre por 7%. A faixa de capacidade, por sua vez, respondeu por "
    "apenas 5% da importância, resultado que reforça o achado da fase anterior: "
    "o indicador oficial de capacidade pesa pouco na decisão quando comparado "
    "ao comportamento real do município, captado pelo grupo. Constata-se, "
    "assim, uma convergência entre as duas fases na relativização do poder "
    "explicativo do rótulo formal, convergência que, por derivar de técnicas "
    "distintas, confere maior solidez ao argumento central do estudo.",
]
for p in res3:
    add_par(p)

add_table(
    ["Variável", "Importância"],
    [
        ["Valor solicitado", "63%"],
        ["Faixa populacional + grupo comportamental", "15%"],
        ["Unidade da federação", "10%"],
        ["Tipo de desastre", "7%"],
        ["Faixa de capacidade (ICM)", "5%"],
    ],
    caption="Tabela 3 - Importância das variáveis no modelo de classificação",
)
add_source("Fonte: elaboração própria.")

add_par(
    "O predomínio do valor solicitado como determinante da aprovação comporta "
    "uma leitura ambivalente. De um lado, indica que o sistema de análise "
    "reage de modo sensível à magnitude dos pedidos, submetendo a maior "
    "escrutínio as solicitações que mais se afastam do padrão, o que é "
    "coerente com a expectativa de um controle atento ao risco. De outro lado, "
    "a centralidade do valor pode sinalizar que outras dimensões relevantes, "
    "como a severidade efetiva do evento ou a qualidade técnica do plano de "
    "trabalho, encontram-se sub-representadas nos dados disponíveis e, por "
    "isso, pesam menos na decisão. Diante dessa ambivalência, o modelo de "
    "classificação deve ser compreendido como um espelho do processo "
    "decisório existente, capaz de revelar suas regularidades, e não como um "
    "padrão normativo do que a decisão deveria ser."
)

add_heading("4.4 Estimativa de “valor justo” e auditoria automatizada",
            level=2)
res4 = [
    "A quarta fase voltou-se à estimativa de faixas de valor esperado e à "
    "sinalização de desvios. A regressão quantílica por gradient boosting "
    "estimou, para cada combinação de atributos, os limites P10, P50 e P90 do "
    "valor solicitado. A cobertura do intervalo entre P10 e P90 alcançou "
    "76,3%, ou seja, cerca de três quartos dos processos situaram-se dentro da "
    "faixa prevista. O coeficiente de determinação no espaço logarítmico, "
    "porém, foi baixo, da ordem de 0,11, o que evidencia a elevada "
    "variabilidade do fenômeno e recomenda que o modelo seja interpretado como "
    "instrumento de triagem, e não como preditor de valor exato. Essa "
    "ressalva, longe de invalidar a abordagem, delimita seu uso adequado: o "
    "modelo é útil para sinalizar afastamentos grosseiros em relação ao "
    "padrão, mas não para estabelecer, com precisão, o valor que cada "
    "reconstrução deveria custar.",

    "A aplicação do modelo a 3.893 processos auditados produziu uma "
    "classificação automática em três categorias. A maior parte dos processos, "
    "correspondente a 66% do total, foi classificada como normal, por situar-"
    "se dentro da faixa esperada. Outros 19% foram classificados como baixo, "
    "indicando possível subdimensionamento, enquanto 15%, equivalentes a 580 "
    "processos, foram classificados como alto, sinalizando valores "
    "potencialmente superdimensionados. Esse último conjunto constitui o "
    "público prioritário para auditoria, pois concentra os casos que mais se "
    "afastam do padrão estimado a partir de processos comparáveis. A categoria "
    "de subdimensionamento, por sua vez, embora menos associada ao risco de "
    "desperdício, merece atenção por sinalizar a possibilidade de planos "
    "incompletos ou de necessidades não plenamente cobertas pelos valores "
    "solicitados.",

    "A categoria de subdimensionamento, embora tenha recebido menos atenção na "
    "literatura sobre auditoria, encerra implicações relevantes para a "
    "efetividade da política. Valores sistematicamente inferiores ao padrão "
    "esperado podem indicar planos de trabalho elaborados sob restrição "
    "técnica, capacidade limitada de estimar custos ou, ainda, subnotificação "
    "de danos por municípios com menor estrutura administrativa. Sob essa "
    "perspectiva, o subdimensionamento não é um problema oposto ao "
    "superfaturamento, mas uma face complementar da mesma fragilidade "
    "institucional, na medida em que ambos remetem à dificuldade de traduzir a "
    "necessidade real em uma solicitação tecnicamente consistente. Diante "
    "disso, a triagem automatizada presta-se não apenas a coibir excessos, mas "
    "também a identificar municípios que poderiam estar deixando de pleitear "
    "recursos compatíveis com os danos efetivamente sofridos.",
]
for p in res4:
    add_par(p)

add_table(
    ["Classificação", "Processos", "%", "Interpretação"],
    [
        ["Normal", "2.556", "66%", "Dentro da faixa esperada"],
        ["Baixo", "757", "19%", "Possível subdimensionamento"],
        ["Alto", "580", "15%", "Valor potencialmente superdimensionado"],
    ],
    caption="Tabela 4 - Resultado da auditoria automatizada dos valores",
)
add_source("Fonte: elaboração própria.")

res4b = [
    "O exame dos dez maiores desvios positivos ilustra a natureza heterogênea "
    "dos casos sinalizados. O município de Nova Monte Verde, no Mato Grosso, "
    "apresentou solicitação de R$ 2,5 bilhões diante de um limite P90 estimado "
    "em cerca de R$ 3,1 milhões, o que corresponde a um desvio de mais de "
    "235.000% e sugere, com elevada probabilidade, erro de registro. Outros "
    "casos, como Jaboatão dos Guararapes, em Pernambuco, e Osasco, em São "
    "Paulo, exibem desvios igualmente expressivos, ainda que de magnitude "
    "inferior. Observa-se que a maioria dos casos de alto risco associa-se a "
    "tempestades convectivas e a enxurradas, e que os processos "
    "correspondentes encontram-se em situações variadas, incluindo "
    "indeferimento, análise e transferência de recurso. Essa heterogeneidade "
    "de situações reforça que a sinalização algorítmica não equivale a um juízo "
    "de irregularidade, mas a uma priorização para verificação humana.",
]
for p in res4b:
    add_par(p)

add_table(
    ["Município (UF)", "Valor solicitado (R$)", "Limite P90 (R$)", "Desvio"],
    [
        ["Nova Monte Verde (MT)", "2.500.000.000", "3.131.467", "235.558%"],
        ["Jaboatão dos Guararapes (PE)", "402.398.352", "6.514.830", "29.867%"],
        ["Rio de Janeiro (RJ)", "172.431.000", "7.202.435", "10.220%"],
        ["Osasco (SP)", "268.376.257", "13.275.506", "9.160%"],
        ["Roca Sales (RS)", "109.436.648", "6.331.731", "9.058%"],
    ],
    caption="Tabela 5 - Cinco maiores desvios positivos identificados pela auditoria",
)
add_source("Fonte: elaboração própria a partir dos dados da auditoria de valores.")

res4c = [
    "O caso de Nova Monte Verde sintetiza tanto a utilidade quanto os limites "
    "da auditoria automatizada. A "
    "magnitude do desvio, da ordem de 235.000%, torna implausível que se trate "
    "de uma solicitação genuína, apontando, com grande probabilidade, para um "
    "erro de digitação no registro do valor. Esse tipo de ocorrência, embora "
    "não configure irregularidade no sentido de má-fé, compromete a qualidade "
    "das bases de dados e pode distorcer estatísticas agregadas, como as "
    "médias por faixa discutidas anteriormente. Dessa maneira, a detecção "
    "automatizada presta um serviço duplo: além de sinalizar possíveis "
    "superfaturamentos, contribui para a depuração dos próprios dados "
    "administrativos, ao expor inconsistências que, de outro modo, "
    "permaneceriam ocultas no volume dos registros.",

    "A quinta fase, de caráter complementar, aplicou a floresta de isolamento "
    "ao conjunto de processos e identificou 320 observações anômalas, "
    "compatíveis, em larga medida, com os casos sinalizados pela regressão "
    "quantílica. A convergência entre técnicas distintas de detecção confere "
    "maior solidez à identificação dos casos prioritários, pois reduz a "
    "probabilidade de que um alerta decorra de particularidade de um único "
    "método. Como desdobramento instrumental, projetou-se um simulador de "
    "alertas, concebido para permitir a consulta interativa das estimativas de "
    "valor esperado e de probabilidade de aprovação, de modo a integrar os "
    "achados das fases anteriores em uma ferramenta de apoio à decisão para os "
    "gestores. O simulador traduz, assim, o conjunto da análise em um artefato "
    "operacional, ainda que sua validação em ambiente real permaneça como "
    "tarefa futura.",
]
for p in res4c:
    add_par(p)

add_heading("4.5 Antecipação da severidade dos desastres", level=2)
res5 = [
    "A solução de previsão revelou os principais vetores que determinam a "
    "gravidade de um desastre no Brasil. A análise da importância das "
    "variáveis apontou o tipo de desastre como fator mais decisivo, com 28% da "
    "importância média, seguido da localização do município, com 22%, da "
    "população, com 18%, do produto interno bruto per capita, com 15%, e da "
    "cobertura de atenção básica em saúde, com 10%. Constata-se, assim, que a "
    "severidade não decorre apenas de fatores socioeconômicos, mas, de modo "
    "preponderante, da natureza e da localização do evento. Esse resultado "
    "dialoga diretamente com os achados da solução de reconstrução, ao "
    "confirmar, por uma via independente, o peso da localização e do tipo de "
    "desastre na determinação tanto dos danos quanto da demanda por recursos.",
]
for p in res5:
    add_par(p)

add_table(
    ["Fator preditivo", "Importância média"],
    [
        ["Tipo de desastre", "28%"],
        ["Localização do município", "22%"],
        ["População do município", "18%"],
        ["PIB per capita", "15%"],
        ["Cobertura de atenção básica", "10%"],
    ],
    caption="Tabela 6 - Principais fatores preditivos da severidade dos desastres",
)
add_source("Fonte: elaboração própria a partir da solução de previsão de severidade.")

add_par(
    "A aplicação prática dessa capacidade preditiva é direta. Ao antecipar a "
    "magnitude provável dos danos de um evento iminente, o instrumento permite "
    "que a defesa civil priorize recursos, que os gestores dimensionem "
    "orçamentos e que as ações de prevenção e preparação sejam direcionadas, "
    "com maior precisão, aos territórios e aos tipos de evento de maior risco. "
    "A disponibilização da solução em ambiente de demonstração de acesso "
    "público reforça a viabilidade de sua incorporação à rotina dos órgãos "
    "responsáveis e ilustra o caráter operacional, e não meramente "
    "exploratório, do conjunto do projeto."
)

# ===========================================================================
add_heading("5 Discussão")
disc = [
    "Os resultados convergem para um argumento central: os indicadores formais "
    "de capacidade institucional, embora úteis, capturam apenas parcialmente o "
    "comportamento dos municípios diante dos desastres. O paradoxo da faixa D, "
    "segundo o qual os municípios de menor capacidade concentram a maior "
    "parcela dos recursos, pode ser interpretado à luz da produção social da "
    "vulnerabilidade (VALENCIO, 2009): a fragilidade institucional, "
    "frequentemente associada a infraestrutura precária e a baixa capacidade "
    "de prevenção, tende a amplificar os danos quando o evento adverso se "
    "materializa. Sob essa leitura, a maior demanda não decorre da capacidade "
    "em si, mas das vulnerabilidades que a baixa capacidade ajuda a perpetuar. "
    "Cabe reiterar, no entanto, que a elevada influência de valores extremos "
    "sobre a média recomenda prudência, pois parte do diferencial observado "
    "pode refletir casos atípicos, inclusive erros de registro, o que impõe "
    "que a interpretação do paradoxo seja sustentada pela análise conjunta da "
    "média e da mediana.",

    "A constatação de que o grupo comportamental supera a faixa de capacidade "
    "na explicação tanto da segmentação quanto da aprovação dos processos tem "
    "implicações práticas relevantes. Verifica-se que a observação do "
    "comportamento efetivo, derivada de dados transacionais, oferece um "
    "substrato informacional mais aderente à realidade do que a classificação "
    "oficial isolada. Diante disso, a focalização de políticas de "
    "fortalecimento institucional e de mecanismos de monitoramento poderia "
    "beneficiar-se da combinação entre o indicador formal e a leitura "
    "comportamental, em vez de apoiar-se exclusivamente no primeiro. Essa "
    "combinação não significa abandonar o índice de capacidade, cuja função de "
    "síntese estrutural permanece válida, mas integrá-lo a uma camada "
    "adicional de informação que registra o comportamento concreto de cada "
    "município ao longo do tempo.",

    "No que tange à auditoria, os resultados indicam que a triagem algorítmica "
    "é viável e potencialmente valiosa para o controle de gastos públicos. A "
    "identificação de 580 processos com valores acima do padrão esperado "
    "permite direcionar a atenção dos órgãos de controle a um subconjunto "
    "gerenciável de casos, em consonância com a lógica de complementaridade "
    "entre controle e capacidade seletiva de fiscalização (SPECK, 2002). "
    "Importa, contudo, delimitar o alcance dessa contribuição. O baixo "
    "coeficiente de determinação do modelo de valor evidencia que as "
    "estimativas comportam incerteza substancial, de modo que a sinalização "
    "deve ser tomada como hipótese a verificar, e não como veredito. A "
    "presença de casos como o de Nova Monte Verde, provável erro de digitação, "
    "ilustra que parcela dos alertas decorre de problemas de qualidade dos "
    "dados, e não de irregularidades substantivas, o que reforça a "
    "indispensabilidade do julgamento humano na etapa de verificação.",

    "Há, ademais, uma dimensão ética e institucional a considerar na adoção de "
    "ferramentas automatizadas no setor público. A sinalização de um processo "
    "como atípico pode produzir efeitos reputacionais sobre gestores e "
    "municípios, ainda que a verificação posterior conclua pela regularidade "
    "do pedido. Por essa razão, a incorporação de tais instrumentos exige "
    "protocolos que assegurem o contraditório, a transparência dos critérios e "
    "a reserva da decisão final ao agente público competente. A "
    "interpretabilidade das técnicas escolhidas neste estudo, ao permitir que "
    "se explicitem os fatores associados a cada sinalização, contribui para "
    "essa exigência, mas não a esgota, pois a legitimidade do uso de algoritmos "
    "na administração depende também de arranjos institucionais que "
    "disciplinem sua aplicação.",

    "Dessa maneira, o estudo apresenta limitações que circunscrevem suas "
    "conclusões. A natureza secundária e administrativa dos dados implica "
    "dependência da qualidade do registro original, e a ausência de variáveis "
    "sobre a severidade física dos eventos restringe a capacidade de isolar o "
    "efeito da vulnerabilidade. Ademais, a definição binária da variável de "
    "aprovação simplifica a diversidade de situações processuais, e a exclusão "
    "dos processos em análise pode introduzir viés de seleção. O baixo poder "
    "explicativo do modelo de valor, por fim, recomenda cautela na sua "
    "utilização para fins distintos da triagem. Essas ressalvas não invalidam "
    "os achados, mas delimitam o terreno em que são válidos e indicam a "
    "necessidade de cautela na sua tradução para decisões de política.",

    "A articulação entre as duas soluções configura um ecossistema que percorre "
    "o ciclo completo de gestão de riscos. O motor de previsão atua a montante, "
    "nas fases de prevenção, mitigação e preparação, ao antecipar a severidade "
    "dos eventos e informar a alocação prévia de recursos; o sistema de "
    "inteligência da reconstrução atua a jusante, na fase de recuperação, ao "
    "otimizar a destinação das verbas e auditar os processos. Tomadas em "
    "conjunto, as ferramentas deslocam a gestão de desastres de um modelo "
    "predominantemente reativo, centrado na resposta ao evento já ocorrido, "
    "para um modelo proativo, fundado na antecipação e no uso contínuo da "
    "informação. Esse deslocamento, contudo, não se realiza apenas pela "
    "disponibilidade técnica dos instrumentos, mas depende de sua integração "
    "aos arranjos institucionais e às rotinas decisórias dos órgãos "
    "competentes, sem a qual a inteligência de dados permanece como potencial "
    "não realizado.",
]
for p in disc:
    add_par(p)

# ===========================================================================
add_heading("6 Considerações finais")
conc = [
    "Este artigo analisou a relação entre a capacidade institucional municipal "
    "e o comportamento de demanda por recursos de reconstrução pós-desastre, "
    "além de avaliar a viabilidade de auditoria automatizada desses gastos, a "
    "partir de 6.385 processos registrados entre 2017 e 2025. Os resultados "
    "evidenciaram que municípios de baixa capacidade concentram a maior "
    "parcela dos recursos e exibem valor médio quase três vezes superior ao "
    "dos municípios de alta capacidade, ainda que esse diferencial seja "
    "fortemente moldado por casos extremos. Demonstrou-se, ademais, que o "
    "comportamento efetivo de demanda é mais informativo que o rótulo oficial "
    "de capacidade, tanto na segmentação dos municípios quanto na explicação "
    "da aprovação dos processos, e que um modelo de triagem foi capaz de "
    "sinalizar 580 processos com valores potencialmente superdimensionados.",

    "A principal contribuição do trabalho reside em articular um diagnóstico "
    "empírico sobre a desigualdade de capacidades a uma proposta instrumental "
    "de governança, fundada em técnicas de aprendizado de máquina concebidas "
    "como apoio, e não como substituto, do julgamento técnico. A triagem "
    "automatizada de valores mostrou-se capaz de priorizar um subconjunto "
    "gerenciável de processos para auditoria, contribuindo para o uso "
    "eficiente dos recursos de fiscalização, ao passo que a segmentação "
    "comportamental ofereceu um mapa mais detalhado da demanda do que a "
    "classificação por faixas. Reconhece-se, todavia, que a incerteza das "
    "estimativas e os problemas de qualidade dos dados impõem limites a essa "
    "aplicação, que deve ser compreendida no quadro de uma governança que "
    "preserve o contraditório e a decisão humana.",

    "Para além das contribuições analíticas, o estudo sugere implicações para o "
    "desenho da política de reconstrução. A constatação de que os recursos se "
    "concentram em municípios de menor capacidade, combinada à evidência de "
    "recorrência territorial dos eventos, indica que a eficácia da política "
    "depende de sua articulação a medidas de fortalecimento institucional e de "
    "prevenção de riscos. Transferir volumes elevados a entes com baixa "
    "capacidade de execução, sem investir simultaneamente na construção dessa "
    "capacidade, tende a reproduzir o ciclo de vulnerabilidade que a política "
    "pretende romper. Nesse sentido, os instrumentos analíticos aqui propostos "
    "podem servir não apenas ao controle dos gastos, mas também à identificação "
    "dos territórios e dos perfis municipais que mais se beneficiariam de "
    "estratégias estruturais, deslocando o foco da resposta reativa para a "
    "gestão antecipada do risco.",

    "Como agenda futura, recomenda-se a incorporação de variáveis sobre a "
    "severidade física dos eventos e sobre as condições socioeconômicas dos "
    "municípios, capazes de qualificar a interpretação do paradoxo da "
    "capacidade; o aprimoramento dos modelos preditivos mediante a inclusão de "
    "novas fontes de dados; e o desenvolvimento e a validação, em ambiente "
    "real, do simulador de alertas, de modo a avaliar sua utilidade efetiva "
    "para os gestores. Seria igualmente proveitoso investigar, por meio de "
    "estudos qualitativos, os casos sinalizados como atípicos, a fim de "
    "discriminar entre erros de registro, situações excepcionais legítimas e "
    "eventuais irregularidades. Tais desdobramentos poderiam consolidar a "
    "integração entre evidência empírica e instrumentos de gestão no campo das "
    "políticas de reconstrução pós-desastre, contribuindo para que o aumento "
    "da capacidade analítica do Estado se converta em maior efetividade e "
    "transparência na aplicação dos recursos públicos.",
]
for p in conc:
    add_par(p)

# ===========================================================================
add_heading("Referências")
refs = [
    "ABRUCIO, Fernando Luiz. A coordenação federativa no Brasil: a experiência "
    "do período FHC e os desafios do governo Lula. Revista de Sociologia e "
    "Política, Curitiba, n. 24, p. 41-67, 2005.",

    "BRASIL. Lei nº 12.608, de 10 de abril de 2012. Institui a Política "
    "Nacional de Proteção e Defesa Civil. Brasília, DF: Presidência da "
    "República, 2012.",

    "BRASIL. Ministério do Meio Ambiente. Plano Nacional de Adaptação à "
    "Mudança do Clima: estratégia setorial de gestão de risco de desastres. "
    "Brasília, DF: MMA, 2016. Disponível em: https://www.gov.br/mma/pt-br/"
    "assuntos/mudanca-do-clima/clima/repositorio/"
    "pna_estrategia_de_gestao_de_risco_de_desastres.pdf. Acesso em: "
    "23 jun. 2026.",

    "BRASIL. Ministério da Saúde. DATASUS: informações de saúde — cobertura da "
    "atenção básica. Brasília, DF: Ministério da Saúde, 2024. Disponível em: "
    "https://datasus.saude.gov.br. Acesso em: 23 jun. 2026.",

    "BREIMAN, Leo. Random forests. Machine Learning, v. 45, n. 1, p. 5-32, 2001.",

    "CENTRO UNIVERSITÁRIO DE ESTUDOS E PESQUISAS SOBRE DESASTRES (CEPED/UFSC). "
    "Atlas digital de desastres no Brasil. Florianópolis: CEPED/UFSC, 2023. "
    "Disponível em: https://atlasdigital.mdr.gov.br. Acesso em: 23 jun. 2026.",

    "CONFEDERAÇÃO NACIONAL DE MUNICÍPIOS (CNM). Panorama dos desastres no "
    "Brasil: 2013 a 2024. Brasília, DF: CNM, 2025. (Estudo Técnico, n. "
    "05/2025). Disponível em: https://cnm.org.br/storage/biblioteca/2025/"
    "Estudos_Tecnicos/ET_DEFCIVIVL_05-2025_Panorama_dos_Desastres_no_Brasil_"
    "2013_a_2024.pdf. Acesso em: 23 jun. 2026.",

    "EQUIPE 5 - INFRAESTRUTURA. Modelo preditivo de severidade de desastres "
    "para a defesa civil: aplicação de demonstração. [S. l.], 2025. Disponível "
    "em: https://github.com/fabiokp/defesa_civil_deploy. Acesso em: 23 jun. 2026.",

    "FRIEDMAN, Jerome H. Greedy function approximation: a gradient boosting "
    "machine. The Annals of Statistics, v. 29, n. 5, p. 1189-1232, 2001.",

    "GRIN, Eduardo José; ABRUCIO, Fernando Luiz. Quão distantes estamos de um "
    "sistema municipal de capacidades estatais no Brasil? Revista de "
    "Administração Pública, Rio de Janeiro, v. 52, n. 4, p. 691-712, 2018.",

    "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). Regiões de "
    "influência das cidades: 2018. Rio de Janeiro: IBGE, 2020.",

    "INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). Produto interno "
    "bruto dos municípios: 2021. Rio de Janeiro: IBGE, 2023. Disponível em: "
    "https://www.ibge.gov.br/estatisticas/economicas/contas-nacionais/"
    "9088-produto-interno-bruto-dos-municipios.html. Acesso em: 23 jun. 2026.",

    "KOENKER, Roger; BASSETT JR., Gilbert. Regression quantiles. Econometrica, "
    "v. 46, n. 1, p. 33-50, 1978.",

    "LIU, Fei Tony; TING, Kai Ming; ZHOU, Zhi-Hua. Isolation forest. In: IEEE "
    "International Conference on Data Mining, 8., 2008. Proceedings. Pisa: "
    "IEEE, 2008. p. 413-422.",

    "MACQUEEN, James. Some methods for classification and analysis of "
    "multivariate observations. In: Berkeley Symposium on Mathematical "
    "Statistics and Probability, 5., 1967. Proceedings. Berkeley: University "
    "of California Press, 1967. p. 281-297.",

    "MARENCO, André; STROHSCHOEN, Maria Tereza Blanco; JONER, William. "
    "Capacidade estatal, burocracia e tributação nos municípios brasileiros. "
    "Revista de Sociologia e Política, Curitiba, v. 25, n. 64, p. 3-21, 2017.",

    "SPECK, Bruno Wilhelm (org.). Caminhos da transparência: análise dos "
    "componentes de um sistema nacional de integridade. Campinas: Editora da "
    "Unicamp, 2002.",

    "VALENCIO, Norma (org.). Sociologia dos desastres: construção, interfaces "
    "e perspectivas no Brasil. São Carlos: RiMa Editora, 2009.",
]
for r in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.line_spacing = 1.0
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(r)
    _set_run_font(run, size=12)

# ---------------------------------------------------------------------------
out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                       "06_relatorios")
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, "artigo_cientifico_enap.docx")
doc.save(out_path)
print("Documento salvo em:", out_path)
